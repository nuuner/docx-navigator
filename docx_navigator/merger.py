"""Core module for merging DOCX files with navigation."""

from pathlib import Path
from collections import defaultdict
from typing import List, Tuple, Optional, Dict, Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docxcompose.composer import Composer


def add_bookmark(paragraph, name: str) -> None:
    """Add a bookmark to a paragraph.
    
    Args:
        paragraph: The paragraph to add the bookmark to
        name: The name of the bookmark
    """
    bid = str(abs(hash(name)) % 2_000_000_000)
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bid)
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bid)
    p = paragraph._p
    p.insert(0, start)
    p.append(end)


def add_internal_link(paragraph, text: str, anchor_name: str) -> None:
    """Add an internal hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the link to
        text: The link text
        anchor_name: The bookmark name to link to
    """
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_name)
    
    run_element = OxmlElement('w:r')
    
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)
    run_element.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    run_element.append(t)
    
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def parse_filename(filename: str, category_separator: str = "_") -> Tuple[str, str]:
    """Parse a filename into category and label.
    
    Args:
        filename: The filename to parse (with or without extension)
        category_separator: The separator between category and label
        
    Returns:
        Tuple of (category, label)
    """
    stem = Path(filename).stem
    if category_separator in stem:
        category, label = stem.split(category_separator, 1)
    else:
        category, label = "General", stem
    return category.strip(), label.strip()


def group_files_by_category(
    files: List[str], 
    category_separator: str = "_"
) -> Dict[str, List[Tuple[str, str]]]:
    """Group files by their category.
    
    Args:
        files: List of file paths
        category_separator: The separator between category and label
        
    Returns:
        Dictionary mapping category to list of (filepath, label) tuples
    """
    groups = defaultdict(list)
    for filepath in files:
        category, label = parse_filename(filepath, category_separator)
        groups[category].append((filepath, label or Path(filepath).stem))
    return dict(groups)


def create_menu_document(
    groups: Dict[str, List[Tuple[str, str]]],
    menu_title: str = "Menu",
    toc_depth: int = 2
) -> Document:
    """Create a document with the navigation menu.
    
    Args:
        groups: Dictionary mapping category to list of (filepath, label) tuples
        menu_title: The title for the menu
        toc_depth: Maximum heading depth for the menu
        
    Returns:
        Document with the menu
    """
    doc = Document()
    
    menu_heading = doc.add_heading(menu_title, level=1)
    add_bookmark(menu_heading, "menu")
    
    for category, items in sorted(groups.items()):
        if toc_depth >= 2:
            cat_heading = doc.add_heading(category, level=2)
            cat_heading_format = cat_heading.paragraph_format
            cat_heading_format.space_before = Pt(6)
            cat_heading_format.space_after = Pt(3)
        
        for filepath, label in items:
            anchor = f"doc_{abs(hash(filepath)) % 1000000000}"
            p = doc.add_paragraph("  ")
            add_internal_link(p, label, anchor)
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)
            p_format.line_spacing = 1.0
    
    doc.add_page_break()
    
    return doc


def merge_documents(
    input_files: List[str],
    output_path: str,
    menu_title: str = "Menu",
    back_label: str = "Back to menu",
    category_separator: str = "_",
    toc_depth: int = 2,
    keep_styles: bool = True,
    dry_run: bool = False
) -> Optional[str]:
    """Merge multiple DOCX files with navigation.
    
    Args:
        input_files: List of input file paths
        output_path: Path for the output file
        menu_title: Title for the navigation menu
        back_label: Label for back-to-menu links
        category_separator: Separator for category parsing
        toc_depth: Maximum heading depth for menu
        keep_styles: Whether to preserve original styles
        dry_run: If True, don't write output file
        
    Returns:
        Path to the output file if successful, None if dry_run
    """
    if not input_files:
        raise ValueError("No input files provided")
    
    groups = group_files_by_category(input_files, category_separator)
    
    if dry_run:
        print(f"Would merge {len(input_files)} files into {output_path}")
        print("\nFile grouping:")
        for category, items in sorted(groups.items()):
            print(f"\n{category}:")
            for filepath, label in items:
                print(f"  - {label} ({Path(filepath).name})")
        return None
    
    base_doc = create_menu_document(groups, menu_title, toc_depth)
    composer = Composer(base_doc)
    
    for filepath in input_files:
        if not Path(filepath).exists():
            print(f"Warning: File not found: {filepath}")
            continue
            
        anchor = f"doc_{abs(hash(filepath)) % 1000000000}"
        
        stub = Document()
        heading = stub.add_heading(Path(filepath).stem, level=1)
        add_bookmark(heading, anchor)
        
        back_paragraph = stub.add_paragraph()
        add_internal_link(back_paragraph, f"⬅ {back_label}", "menu")
        stub.add_paragraph()
        
        composer.append(stub)
        
        try:
            doc_to_append = Document(filepath)
            if keep_styles:
                composer.append(doc_to_append)
            else:
                composer.append(doc_to_append, preserve_styles=False)
        except Exception as e:
            print(f"Error appending {filepath}: {e}")
            continue
        
        page_break_doc = Document()
        page_break_doc.add_page_break()
        composer.append(page_break_doc)
    
    composer.save(output_path)
    return output_path