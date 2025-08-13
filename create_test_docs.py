#!/usr/bin/env python3
"""Create sample DOCX files for testing."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_doc(filename: str, title: str, content: list):
    """Create a sample DOCX file with given content."""
    doc = Document()
    
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for item in content:
        if isinstance(item, dict):
            if item['type'] == 'heading':
                doc.add_heading(item['text'], level=item.get('level', 2))
            elif item['type'] == 'paragraph':
                p = doc.add_paragraph(item['text'])
                if item.get('style'):
                    p.style = item['style']
            elif item['type'] == 'bullet':
                doc.add_paragraph(item['text'], style='List Bullet')
            elif item['type'] == 'page_break':
                doc.add_page_break()
        else:
            doc.add_paragraph(item)
    
    doc.save(filename)
    print(f"Created: {filename}")


def main():
    """Create sample test documents."""
    
    create_sample_doc(
        "Finance_Quarterly Report Q1.docx",
        "Q1 Financial Report",
        [
            {"type": "heading", "text": "Executive Summary", "level": 2},
            {"type": "paragraph", "text": "This quarter showed strong growth across all key metrics."},
            {"type": "heading", "text": "Revenue", "level": 2},
            {"type": "bullet", "text": "Total Revenue: $5.2M"},
            {"type": "bullet", "text": "Growth: 15% YoY"},
            {"type": "bullet", "text": "New Customers: 127"},
            {"type": "heading", "text": "Expenses", "level": 2},
            {"type": "paragraph", "text": "Operating expenses remained within budget at $3.1M."},
        ]
    )
    
    create_sample_doc(
        "Finance_Quarterly Report Q2.docx",
        "Q2 Financial Report",
        [
            {"type": "heading", "text": "Executive Summary", "level": 2},
            {"type": "paragraph", "text": "Q2 continued the positive momentum from Q1."},
            {"type": "heading", "text": "Revenue", "level": 2},
            {"type": "bullet", "text": "Total Revenue: $6.1M"},
            {"type": "bullet", "text": "Growth: 17% YoY"},
            {"type": "bullet", "text": "New Customers: 156"},
            {"type": "heading", "text": "Outlook", "level": 2},
            {"type": "paragraph", "text": "We expect continued growth in Q3 and Q4."},
        ]
    )
    
    create_sample_doc(
        "HR_Employee Handbook.docx",
        "Employee Handbook",
        [
            {"type": "heading", "text": "Welcome", "level": 2},
            {"type": "paragraph", "text": "Welcome to our company! This handbook contains important information."},
            {"type": "heading", "text": "Company Values", "level": 2},
            {"type": "bullet", "text": "Integrity"},
            {"type": "bullet", "text": "Innovation"},
            {"type": "bullet", "text": "Collaboration"},
            {"type": "heading", "text": "Policies", "level": 2},
            {"type": "paragraph", "text": "All employees must follow company policies as outlined below."},
            {"type": "heading", "text": "Time Off", "level": 3},
            {"type": "paragraph", "text": "Employees receive 15 days of PTO annually."},
        ]
    )
    
    create_sample_doc(
        "HR_Payroll Guidelines.docx",
        "Payroll Guidelines",
        [
            {"type": "heading", "text": "Pay Schedule", "level": 2},
            {"type": "paragraph", "text": "Employees are paid bi-weekly on Fridays."},
            {"type": "heading", "text": "Direct Deposit", "level": 2},
            {"type": "paragraph", "text": "All employees must set up direct deposit within 30 days."},
            {"type": "heading", "text": "Benefits", "level": 2},
            {"type": "bullet", "text": "Health Insurance"},
            {"type": "bullet", "text": "401(k) Matching"},
            {"type": "bullet", "text": "Life Insurance"},
        ]
    )
    
    create_sample_doc(
        "Marketing_Brand Guidelines.docx",
        "Brand Guidelines",
        [
            {"type": "heading", "text": "Brand Identity", "level": 2},
            {"type": "paragraph", "text": "Our brand represents innovation and reliability."},
            {"type": "heading", "text": "Logo Usage", "level": 2},
            {"type": "paragraph", "text": "The logo must maintain minimum clear space of 0.5 inches."},
            {"type": "heading", "text": "Color Palette", "level": 2},
            {"type": "bullet", "text": "Primary: Blue (#0066CC)"},
            {"type": "bullet", "text": "Secondary: Gray (#666666)"},
            {"type": "bullet", "text": "Accent: Green (#00AA44)"},
        ]
    )
    
    create_sample_doc(
        "Marketing_Campaign Plan 2025.docx",
        "2025 Marketing Campaign",
        [
            {"type": "heading", "text": "Campaign Overview", "level": 2},
            {"type": "paragraph", "text": "The 2025 campaign focuses on digital transformation."},
            {"type": "heading", "text": "Target Audience", "level": 2},
            {"type": "bullet", "text": "Enterprise customers"},
            {"type": "bullet", "text": "SMB segment"},
            {"type": "bullet", "text": "Startups"},
            {"type": "heading", "text": "Channels", "level": 2},
            {"type": "paragraph", "text": "We will leverage multiple channels including social media, email, and events."},
            {"type": "heading", "text": "Budget", "level": 2},
            {"type": "paragraph", "text": "Total budget: $2.5M allocated across all channels."},
        ]
    )
    
    print("\nAll test documents created successfully!")


if __name__ == "__main__":
    main()